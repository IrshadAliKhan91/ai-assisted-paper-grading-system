import re
import io
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(doc):
    """Yield paragraphs and tables in true document order.

    python-docx exposes doc.paragraphs and doc.tables as separate lists, which
    loses the interleaving of questions (paragraphs) and their 'Answer:' cells
    (tables). Walking the body element preserves order so each answer cell can
    be tied to the question immediately above it.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def _is_answer_table(table):
    """A 1x1 'Answer:' box that follows a question."""
    if len(table.rows) == 1 and len(table.columns) == 1:
        return table.rows[0].cells[0].text.strip().lower().startswith('answer')
    return False


def _answer_text(table):
    """Extract whatever the teacher wrote in an 'Answer:' cell (label stripped)."""
    raw = table.rows[0].cells[0].text.strip()
    return re.sub(r'^answer\s*:?\s*', '', raw, flags=re.IGNORECASE).strip()


def _read_info_table(table, info):
    """Pull Subject / Total Marks from a 2-column key/value table into `info`."""
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2 or not cells[1]:
            continue
        key = cells[0].lower().rstrip(':').strip()
        if key == 'subject':
            info.setdefault('subject', cells[1])
        elif key in ('total marks', 'total mark', 'max marks', 'maximum marks'):
            m = re.search(r'\d+(?:\.\d+)?', cells[1])
            if m:
                info.setdefault('total_marks', float(m.group(0)))


def parse_template(file_bytes: bytes) -> dict:
    """
    Parse an answer-sheet template (.docx) into an answer key.

    Follows the IM School System layout, reading EVERYTHING dynamically from the
    document (nothing fixed in code):
      - Subject / Total Marks      → the info table (else a heading line).
      - Per-question max_marks      → explicit '[X Marks]' tag
                                      → 'Each question carries X marks' line
                                      → Total Marks / number of questions
                                      → 5.0 (last-resort default only).
      - Model answer per question   → read from the question's 'Answer:' cell if
                                      the teacher filled it; otherwise left empty
                                      to be supplied later by the teacher / AI.

    Returns: { 'subject': str, 'total_marks': float, 'questions': [
        { 'question': str, 'max_marks': float, 'answer': str }, ... ] }
    """
    doc = Document(io.BytesIO(file_bytes))

    q_pattern = re.compile(r'^Q\d+[\.\)\:]?\s*(.*?)(?:\s*\[(\d+(?:\.\d+)?)\s*[mM]arks?\])?\s*$')
    instruction_re = re.compile(r'each\s+question\s+carries\s+(\d+(?:\.\d+)?)\s*marks?', re.IGNORECASE)

    info = {}                 # subject / total_marks from the info table
    heading_candidates = []   # non-question paragraphs (for subject fallback)
    instruction_marks = None
    raw_questions = []        # {'question', 'explicit', 'answer'}
    current = None            # question awaiting its answer cell

    # Single ordered pass so answer cells bind to the question above them.
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if instruction_marks is None:
                m = instruction_re.search(text)
                if m:
                    instruction_marks = float(m.group(1))
            qm = q_pattern.match(text)
            if qm:
                qtext = qm.group(1).strip()
                qtext = re.sub(r'[\.\s\?]+\s*$', '?', qtext)
                if not qtext.endswith('?'):
                    qtext += '?'
                current = {
                    'question': qtext,
                    'explicit': float(qm.group(2)) if qm.group(2) else None,
                    'answer': '',
                }
                raw_questions.append(current)
            else:
                heading_candidates.append(text)
        elif isinstance(block, Table):
            if _is_answer_table(block):
                if current is not None:
                    current['answer'] = _answer_text(block)  # '' if left blank
            else:
                _read_info_table(block, info)

    # ── Subject ──
    subject = info.get('subject')
    if not subject:
        for text in heading_candidates:
            if "SCHOOL SYSTEM" in text.upper() or "INSTRUCTION" in text.upper():
                continue
            if "-" in text and any(k in text for k in ("Test", "Exam", "Paper")):
                subject = text.split("-")[0].strip()
            else:
                subject = text
            break
    if not subject:
        subject = "General Subject"

    # ── Marks (all dynamic) ──
    total_marks = info.get('total_marks')
    num_q = len(raw_questions)
    even_split = (total_marks / num_q) if (total_marks and num_q) else None
    default_mark = instruction_marks or even_split or 5.0

    questions = [{
        "question": q['question'],
        "max_marks": float(q['explicit'] if q['explicit'] is not None else default_mark),
        "answer": q['answer'],
    } for q in raw_questions]

    if total_marks is None:
        total_marks = sum(q["max_marks"] for q in questions)

    return {
        "subject": subject,
        "total_marks": float(total_marks),
        "questions": questions,
    }
